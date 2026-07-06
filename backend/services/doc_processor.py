"""Document processing: parse .docx and HTML into structured elements, split by H2."""
import os
import re
import base64
import io
from typing import Optional

_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# Flag added to a section whose source had verbatim-duplicated content that we
# auto-collapsed during parsing (see _append_element). Surfaced in SectionViewer.
DUP_COLLAPSED_FLAG = "Source had duplicated content (auto-collapsed)"


def iter_block_items(parent):
    """Yield Paragraph and Table objects from a docx body in TRUE document order.

    Unlike ``doc.paragraphs`` / ``doc.tables`` (which are flat and separate, so a
    mid-content table ends up after every paragraph), this walks the body's
    children in order. Content controls (``w:sdt``) are unwrapped so their inner
    paragraphs/tables are yielded once, in place — content ``doc.paragraphs``
    would otherwise drop entirely. Nested content controls are handled by
    recursion; each paragraph/table is yielded exactly once (no duplication).
    """
    from docx.document import Document as _Document
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    part_owner = parent
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent

    def walk(elm):
        for child in elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, part_owner)
            elif isinstance(child, CT_Tbl):
                yield Table(child, part_owner)
            elif child.tag == f'{_W}sdt':
                content = child.find(f'{_W}sdtContent')
                if content is not None:
                    yield from walk(content)

    yield from walk(parent_elm)


def _append_element(elements: list, elem: dict, seen: set, current_h2: Optional[dict]) -> bool:
    """Append ``elem`` to the section verbatim — the document is read as-is.

    Dedup DISABLED: a previous version collapsed a paragraph/list_item whose
    normalized text matched one already seen in the same H2 section. That was too
    aggressive — a sub-bullet like "Presentation" legitimately repeats under
    different parent bullets (Disease 1 / Disease 2), and the dedup destroyed the
    second one. We now reproduce the source faithfully and never drop content.

    Signature (seen / current_h2) is kept so callers don't change; to re-enable a
    SAFER dedup later, key on (heading_context, parent-bullet, text), not text
    alone. ``dup_collapsed_flag`` stays a no-op while this is disabled.
    """
    elements.append(elem)
    return True


def dup_collapsed_flag(elems: list[dict]) -> Optional[str]:
    """Build the reviewer-facing dedup flag for a section's elements, listing the
    lines that were auto-collapsed. Returns None if nothing was collapsed.

    e.g. 'Source had duplicated content (auto-collapsed): "Supportive care", "None"'
    """
    lines: list[str] = []
    for e in elems:
        for line in (e.get("_dup_collapsed_lines") or []):
            if line not in lines:
                lines.append(line)
    if not any(e.get("_dup_collapsed") for e in elems):
        return None
    if not lines:
        return DUP_COLLAPSED_FLAG
    shown = lines[:10]
    parts = ", ".join(f'"{(l[:80] + "…") if len(l) > 80 else l}"' for l in shown)
    if len(lines) > len(shown):
        parts += f", +{len(lines) - len(shown)} more"
    return f"{DUP_COLLAPSED_FLAG}: {parts}"


def _all_runs(para):
    """All runs in a paragraph in document order, INCLUDING runs wrapped in
    inline content controls (<w:sdt>). Google Docs exports wrap suggested/edited
    text in inline SDTs; python-docx's para.runs / para.text skip those and
    silently drop the text. Walking descendant <w:r> elements recovers it."""
    from docx.text.run import Run
    return [Run(r, para) for r in para._element.iter(f'{_W}r')]


def parse_docx(filepath: str) -> list[dict]:
    """Parse a .docx file into a list of element dicts.

    Each element has:
    - type: paragraph, heading, table, image
    - text: plain text content
    - html: HTML representation
    - level: heading level (1-6) for headings, None otherwise
    - heading_context: breadcrumb of parent headings (for nested content)
    - data_uri: base64 data URI for images
    - alt_text: alt text hint for images
    """
    import docx
    from docx.table import Table

    doc = docx.Document(filepath)
    elements = []
    current_headings = {}  # level -> heading text
    seen: set = set()           # normalized content keys within the current H2 section
    current_h2: Optional[dict] = None  # ref to the current H2 heading element (for dup flag)

    # Single pass in true document order: paragraphs and tables interleaved, with
    # content controls unwrapped (see iter_block_items).
    for block in iter_block_items(doc):
        # ── Tables — emit inline at their real position ──────────────────────
        if isinstance(block, Table):
            rows = []
            for row in block.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(cells)
            if rows:
                elements.append({
                    "type": "table",
                    "text": _table_to_text(rows),
                    "html": _table_to_html(rows),
                    "rows": rows,
                    "heading_context": _build_heading_context(current_headings),
                })
            continue

        para = block
        # Edited/exported docs (Pages, some Word round-trips) can reference a
        # missing style — para.style is then None and .name would crash the parse.
        try:
            style_name = (para.style.name or "").lower()
        except AttributeError:
            style_name = ""
        _runs = _all_runs(para)
        text = ''.join(r.text for r in _runs).strip()

        if not text and not _runs:
            continue

        # Check for images in paragraph runs
        _WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        for run in _runs:
            if run._element.findall(f'.//{_W}drawing'):
                # Extract inline images
                for drawing in run._element.findall(f'.//{_W}drawing'):
                    # Read alt text from wp:docPr descr attribute
                    doc_pr = drawing.find(f'.//{{{_WP_NS}}}docPr')
                    alt_text = None
                    if doc_pr is not None:
                        raw_descr = doc_pr.get('descr', '').strip()
                        if raw_descr:
                            alt_text = raw_descr

                    for blip in drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed:
                            try:
                                rel = para.part.rels[embed]
                                image_data = rel.target_part.blob
                                content_type = rel.target_part.content_type or "image/png"
                                b64 = base64.b64encode(image_data).decode('utf-8')
                                data_uri = f"data:{content_type};base64,{b64}"
                                elements.append({
                                    "type": "image",
                                    "text": "",
                                    "html": "",
                                    "data_uri": data_uri,
                                    "alt_text": alt_text,
                                    "heading_context": _build_heading_context(current_headings),
                                })
                            except (KeyError, AttributeError, ValueError):
                                # ValueError: externally-linked image (target_part is
                                # undefined for external rels) — skip, don't crash.
                                pass

        if not text:
            continue

        # Detect heading level
        level = None
        if "heading" in style_name:
            match = re.search(r'(\d)', style_name)
            if match:
                level = int(match.group(1))

        if level is not None:
            # Update heading context
            current_headings[level] = text
            # Clear lower-level headings
            for l in list(current_headings.keys()):
                if l > level:
                    del current_headings[l]

            heading_elem = {
                "type": "heading",
                "text": text,
                "html": f"<h{level}>{text}</h{level}>",
                "level": level,
                "heading_context": _build_heading_context(current_headings),
            }
            elements.append(heading_elem)
            # Each H2 starts a new section — reset the dedup window and track it.
            if level == 2:
                seen = set()
                current_h2 = heading_elem
        else:
            # Build HTML with basic formatting from runs
            html_parts = []
            for run in _runs:
                t = run.text
                if not t:
                    continue
                if run.bold:
                    t = f"<b>{t}</b>"
                if run.italic:
                    t = f"<i>{t}</i>"
                if run.underline:
                    t = f"<u>{t}</u>"
                html_parts.append(t)
            html = "".join(html_parts) or text
            # Word soft line breaks (Shift+Enter) come through python-docx as "\n"
            # in the run/paragraph text. Make them explicit <br> so the structure
            # survives into content_html — the section modal then renders the real
            # lines, and the source we send the AI keeps them on separate lines
            # (instead of collapsing into one run-together block). Mirrors what a
            # copy-paste from Word into a chat would preserve.
            html = html.replace("\n", "<br>")

            # Detect list items by style name or numbering XML
            is_list = False
            list_type = "ul"  # default unordered
            indent_level = 0

            if "list" in style_name:
                is_list = True
                if "number" in style_name or "ordered" in style_name:
                    list_type = "ol"
            else:
                # Check for Word numbering (numPr element)
                num_pr = para._element.find(f'.//{_W}numPr')
                if num_pr is not None:
                    is_list = True
                    # Check numId to distinguish bullet vs number
                    num_id_el = num_pr.find(f'{_W}numId')
                    if num_id_el is not None:
                        num_id_val = int(num_id_el.get(f'{_W}val', '0'))
                        # numId 0 means no numbering
                        if num_id_val == 0:
                            is_list = False

            if is_list:
                # Get indentation level from ilvl
                ilvl_el = para._element.find(f'.//{_W}ilvl')
                if ilvl_el is not None:
                    indent_level = int(ilvl_el.get(f'{_W}val', '0'))
                else:
                    # Try indentation from paragraph properties
                    ind_el = para._element.find(f'.//{_W}ind')
                    if ind_el is not None:
                        left = ind_el.get(f'{_W}left', '0')
                        indent_level = max(0, int(left) // 720)  # 720 twips ≈ 0.5 inch

            if is_list:
                _append_element(elements, {
                    "type": "list_item",
                    "text": text,
                    "html": html,
                    "list_type": list_type,
                    "indent_level": indent_level,
                    "heading_context": _build_heading_context(current_headings),
                }, seen, current_h2)
            else:
                # Check for indentation on regular paragraphs
                ind_el = para._element.find(f'.//{_W}ind')
                if ind_el is not None:
                    left = ind_el.get(f'{_W}left', '0')
                    indent_level = max(0, int(left) // 720)

                if indent_level > 0:
                    _append_element(elements, {
                        "type": "paragraph",
                        "text": text,
                        "html": f'<p style="margin-left:{indent_level * 1.5}em">{html}</p>',
                        "heading_context": _build_heading_context(current_headings),
                    }, seen, current_h2)
                else:
                    _append_element(elements, {
                        "type": "paragraph",
                        "text": text,
                        "html": f"<p>{html}</p>",
                        "heading_context": _build_heading_context(current_headings),
                    }, seen, current_h2)

    return elements


def parse_html(filepath: str) -> list[dict]:
    """Parse an HTML file into structured elements."""
    from bs4 import BeautifulSoup

    with open(filepath, "r", encoding="utf-8") as f:
        html_content = f.read()

    soup = BeautifulSoup(html_content, "html.parser")
    elements = []
    current_headings = {}

    def _li_direct(li, *, as_html: bool) -> str:
        """A list item's OWN content, excluding nested <ul>/<ol> children.
        Without this, get_text() on a parent <li> swallows all descendant items'
        text (and the nested items are also emitted separately) — duplicating and
        flattening nested lists on a copy→paste round-trip."""
        parts = []
        for child in li.children:
            name = getattr(child, "name", None)
            if name in ("ul", "ol"):
                continue
            parts.append(str(child) if (as_html and name) else
                         (child.get_text() if name else str(child)))
        out = "".join(parts)
        return out.strip() if as_html else re.sub(r"\s+", " ", out).strip()

    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table', 'div', 'li']):
        tag = elem.name.lower()

        # List items: emit each with its OWN text + nesting depth + list type so
        # build_content_html can reconstruct nested <ul>/<ol>. Process every <li>
        # (including nested) exactly once; the parent never includes child text.
        if tag == 'li':
            text = _li_direct(elem, as_html=False)
            if not text:
                continue
            list_ancestors = [p for p in elem.parents if getattr(p, "name", None) in ("ul", "ol")]
            indent_level = max(0, len(list_ancestors) - 1)
            list_type = list_ancestors[0].name if list_ancestors else "ul"
            elements.append({
                "type": "list_item",
                "text": text,
                "html": _li_direct(elem, as_html=True) or text,
                "indent_level": indent_level,
                "list_type": list_type,
                "heading_context": _build_heading_context(current_headings),
            })
            continue

        text = elem.get_text(strip=True)
        if not text:
            continue

        if tag.startswith('h') and len(tag) == 2 and tag[1].isdigit():
            level = int(tag[1])
            current_headings[level] = text
            for l in list(current_headings.keys()):
                if l > level:
                    del current_headings[l]
            elements.append({
                "type": "heading",
                "text": text,
                "html": str(elem),
                "level": level,
                "heading_context": _build_heading_context(current_headings),
            })
        elif tag == 'table':
            rows = []
            for tr in elem.find_all('tr'):
                cells = [td.get_text(strip=True) for td in tr.find_all(['td', 'th'])]
                rows.append(cells)
            if rows:
                elements.append({
                    "type": "table",
                    "text": _table_to_text(rows),
                    "html": str(elem),
                    "rows": rows,
                    "heading_context": _build_heading_context(current_headings),
                })
        elif tag in ('p', 'div'):
            # Skip paragraphs/divs nested inside a list item — their text is
            # already captured by the <li> above (avoids duplication).
            if elem.find_parent('li'):
                continue
            elements.append({
                "type": "paragraph",
                "text": text,
                "html": str(elem),
                "heading_context": _build_heading_context(current_headings),
            })

    return elements


def split_by_h2(elements: list[dict]) -> list[tuple[str, list[dict]]]:
    """Split elements into groups by H2 headings.

    Returns list of (heading_text, elements_in_section).
    Content before the first H2 goes into a "Preamble" section.
    """
    sections = []
    current_heading = "Preamble"
    current_elements = []

    for elem in elements:
        if elem.get("type") == "heading" and elem.get("level") == 2:
            if current_elements:
                sections.append((current_heading, current_elements))
            current_heading = elem.get("text", "Untitled Section")
            current_elements = [elem]
        else:
            current_elements.append(elem)

    if current_elements:
        sections.append((current_heading, current_elements))

    return sections


def build_heading_tree(elements: list[dict]) -> list[dict]:
    """Build a nested heading tree (H3/H4) from a section's elements.

    Returns a list of dicts like:
    [{"heading": "Diagnosis", "level": 3, "children": [{"heading": "BNP Levels", "level": 4, "children": []}]}]
    """
    tree = []
    stack = []  # (level, node)

    for elem in elements:
        if elem.get("type") != "heading":
            continue
        level = elem.get("level", 3)
        if level < 3:
            continue  # skip H1/H2 — those are section boundaries

        node = {"heading": elem["text"], "level": level, "children": []}

        # Pop items from stack that are at same or deeper level
        while stack and stack[-1][0] >= level:
            stack.pop()

        if stack:
            stack[-1][1]["children"].append(node)
        else:
            tree.append(node)

        stack.append((level, node))

    return tree


def parse_heading_outline(elements: list[dict]) -> list[dict]:
    """Nested H1–H4 outline from a flat element list. Each heading node is
    {"hid": int, "level": int, "text": str, "children": [...]}. `hid` is the
    heading's 0-based index in document order — the SAME counter
    attach_content_to_curriculum uses, so the aligner can key decisions by hid.

    Headings attach to the nearest previous heading of a shallower level (so a
    skipped level just nests under whatever is open)."""
    roots: list[dict] = []
    stack: list[dict] = []  # open ancestor heading nodes, increasing level
    hid = 0
    for elem in elements:
        if elem.get("type") != "heading":
            continue
        level = elem.get("level", 1)
        node = {"hid": hid, "level": level, "text": elem.get("text", ""), "children": []}
        hid += 1
        while stack and stack[-1]["level"] >= level:
            stack.pop()
        if stack:
            stack[-1]["children"].append(node)
        else:
            roots.append(node)
        stack.append(node)
    return roots


def compare_heading_trees(existing: list[dict], new: list[dict]) -> dict:
    """Compare two heading trees and find matches, new headings, and missing headings.

    Returns:
    {
        "matched": [{"existing": heading, "new": heading}],
        "new": [heading],
        "missing": [heading],
    }
    """
    existing_headings = _flatten_headings(existing)
    new_headings = _flatten_headings(new)

    existing_set = set(existing_headings)
    new_set = set(new_headings)

    matched = [{"existing": h, "new": h} for h in existing_set & new_set]
    added = list(new_set - existing_set)
    missing = list(existing_set - new_set)

    return {"matched": matched, "new": added, "missing": missing}


def _flatten_headings(tree: list[dict]) -> list[str]:
    """Flatten a heading tree into a list of heading strings."""
    headings = []
    for node in tree:
        headings.append(node["heading"])
        if node.get("children"):
            headings.extend(_flatten_headings(node["children"]))
    return headings


def _build_heading_context(headings: dict) -> Optional[str]:
    """Build a heading context string from current heading hierarchy."""
    if not headings:
        return None
    parts = []
    for level in sorted(headings.keys()):
        parts.append(f"H{level}: {headings[level]}")
    return " > ".join(parts)


def build_content_html(elements: list[dict]) -> str:
    """Build HTML from elements, properly grouping list items into <ul>/<ol> tags."""
    html_parts = []
    list_stack = []  # stack of (list_type, indent_level)
    img_counter = 0

    for elem in elements:
        elem_type = elem.get("type", "paragraph")

        if elem_type == "list_item":
            target_indent = elem.get("indent_level", 0)
            lt = elem.get("list_type", "ul")

            # Close deeper lists
            while list_stack and list_stack[-1][1] > target_indent:
                html_parts.append(f"</li></{list_stack[-1][0]}>")
                list_stack.pop()

            if not list_stack or list_stack[-1][1] < target_indent:
                # Open new nested list
                html_parts.append(f"<{lt}>")
                list_stack.append((lt, target_indent))
            else:
                # Same level — close previous <li>
                html_parts.append("</li>")

            html_parts.append(f"<li>{elem.get('html', elem.get('text', ''))}")
        else:
            # Close all open lists before non-list content
            while list_stack:
                html_parts.append(f"</li></{list_stack[-1][0]}>")
                list_stack.pop()

            if elem_type == "image":
                img_counter += 1
                html_parts.append(
                    f'<div class="image-placeholder" data-img-index="{img_counter}">'
                    f'[Image {img_counter}]</div>'
                )
            else:
                html_parts.append(elem.get("html", elem.get("text", "")))

    # Close remaining open lists
    while list_stack:
        html_parts.append(f"</li></{list_stack[-1][0]}>")
        list_stack.pop()

    return "\n".join(html_parts)


def _table_to_text(rows: list[list[str]]) -> str:
    """Convert table rows to plain text."""
    lines = []
    for row in rows:
        lines.append(" | ".join(row))
    return "\n".join(lines)


def _table_to_html(rows: list[list[str]]) -> str:
    """Convert table rows to HTML table."""
    html_parts = ["<table>"]
    for i, row in enumerate(rows):
        html_parts.append("<tr>")
        tag = "th" if i == 0 else "td"
        for cell in row:
            html_parts.append(f"<{tag}>{cell}</{tag}>")
        html_parts.append("</tr>")
    html_parts.append("</table>")
    return "".join(html_parts)


def is_title_junk(text: str, doc_name: str = "", topic_name: str = "") -> bool:
    """True for document-title lines that shouldn't become content: PAEA
    blueprint banners, or lines that are just the document/main-topic name."""
    def norm(s: str) -> str:
        return re.sub(r"\s+", " ", (s or "")).strip().lower()

    t = norm(text)
    if not t:
        return True
    if "paea blueprint" in t:
        return True
    if "blueprint" in t and "effective" in t:
        return True
    # "Effective July 2026" style banner — only when the month is explicit and
    # the line is short/title-like, so prose ("shown effective in 2019 trials")
    # is never dropped.
    if len(t) < 80 and re.search(
        r"effective\s+(january|february|march|april|may|june|july|august|"
        r"september|october|november|december)\s*,?\s*\d{4}", t,
    ):
        return True
    doc = norm(os.path.splitext(doc_name or "")[0])
    topic = norm(topic_name)
    if t and (t == doc or t == topic):
        return True
    if len(t) > 8 and ((doc and t in doc) or (topic and t in topic)):
        return True
    return False


def attach_content_to_curriculum(elements: list, resolution: dict,
                                 main_topic_id: int) -> list:
    """Group elements by the deepest matched curriculum node along each block's
    heading ancestry (rolling up; pre-heading content → main_topic_id). Returns
    [{"node_id": int, "elements": [...]}], groups in first-seen order, elements
    within a group in original document order (so build_content_html image
    numbering and SectionImage.position stay aligned).

    A MATCHED heading (resolution[hid] is a real node) defines a section title and
    is NOT emitted into the body; an UNMATCHED heading (None) is kept as
    sub-structure. A matched node with no content produces no group."""
    groups: dict = {}
    order: list = []
    stack: list = []  # (level, hid) of open headings
    hid = 0

    def emit(node_id: int, elem: dict) -> None:
        if node_id not in groups:
            groups[node_id] = []
            order.append(node_id)
        groups[node_id].append(elem)

    for elem in elements:
        if elem.get("type") == "heading":
            level = elem.get("level", 1)
            while stack and stack[-1][0] >= level:
                stack.pop()
            cur_hid = hid
            stack.append((level, cur_hid))
            hid += 1
            if resolution.get(cur_hid) is None:
                emit(_resolve_node(stack, resolution, main_topic_id), elem)
        else:
            emit(_resolve_node(stack, resolution, main_topic_id), elem)

    return [{"node_id": nid, "elements": groups[nid]} for nid in order]


def _resolve_node(stack: list, resolution: dict, main_topic_id: int) -> int:
    for _level, h in reversed(stack):
        nid = resolution.get(h)
        if nid is not None:
            return nid
    return main_topic_id
