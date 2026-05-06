"""AI-powered heading detection for documents without proper heading styles.

Extracts paragraph metadata from a .docx (text, bold, font size, style name,
indentation) and asks Claude to identify which paragraphs are headings and at
what level (H1–H4). Returns the parsed elements with headings injected.
"""
import json
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Max paragraphs to send for analysis — beyond this we risk token limits
MAX_PARAGRAPHS = 400


def extract_paragraph_hints(filepath: str) -> list[dict]:
    """Extract lightweight paragraph metadata from a .docx for heading analysis.

    Returns list of dicts with:
      index, text, style_name, is_bold, font_size_pt, indent_twips, char_count
    """
    import docx

    doc = docx.Document(filepath)
    hints = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "Normal").lower()

        # Bold: paragraph-level bold or all runs bold
        para_bold = para.runs and all(r.bold for r in para.runs if r.text.strip())

        # Font size: largest font size in runs (pt)
        font_size = None
        for run in para.runs:
            if run.font.size:
                pt = run.font.size.pt
                if font_size is None or pt > font_size:
                    font_size = pt

        # Indentation (left, in twips)
        indent_twips = 0
        pPr = para._element.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
        if pPr is not None:
            ind = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ind')
            if ind is not None:
                left = ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}left', '0')
                try:
                    indent_twips = int(left)
                except ValueError:
                    pass

        hints.append({
            "index": i,
            "text": text[:120],  # truncate to keep tokens low
            "style": style_name,
            "bold": para_bold,
            "font_size": font_size,
            "indent": indent_twips,
            "chars": len(text),
        })

        if len(hints) >= MAX_PARAGRAPHS:
            break

    return hints


def _load_curriculum_topics(version: str = 'v1') -> list[str]:
    """Load flat list of curriculum topic names from the database."""
    try:
        from backend.db import SessionLocal
        from backend.models import Curriculum
        db = SessionLocal()
        try:
            nodes = db.query(Curriculum).filter(Curriculum.version == version).all()
            return [n.name for n in nodes]
        finally:
            db.close()
    except Exception as e:
        logger.warning("Could not load curriculum topics: %s", e)
        return []


def detect_headings_with_ai(
    filepath: str,
    model: str = "claude-haiku-4-5-20251001",
    curriculum_version: str = 'v1',
) -> list[dict]:
    """Call Claude to identify heading paragraphs using curriculum topics as a guide.

    Returns list of {index, level} for paragraphs identified as headings.
    """
    import anthropic
    from backend.config import ANTHROPIC_API_KEY

    hints = extract_paragraph_hints(filepath)
    if not hints:
        return []

    # Load curriculum topics to guide matching
    curriculum_topics = _load_curriculum_topics(curriculum_version)

    # Build a compact numbered list for Claude
    lines = []
    for h in hints:
        parts = [f'{h["index"]}. "{h["text"]}"']
        clues = []
        if h["bold"]:
            clues.append("bold")
        if h["font_size"] and h["font_size"] >= 12:
            clues.append(f"{h['font_size']:.0f}pt")
        if h["indent"] == 0 and h["chars"] < 80:
            clues.append("short+flush")
        if h["style"] not in ("normal", "default paragraph font", ""):
            clues.append(f'style:"{h["style"]}"')
        if clues:
            parts.append(f"[{', '.join(clues)}]")
        lines.append(" ".join(parts))

    paragraph_list = "\n".join(lines)

    curriculum_section = ""
    if curriculum_topics:
        # Show a sample of curriculum topic names to guide Claude
        sample = curriculum_topics[:120]  # keep prompt manageable
        curriculum_section = f"""
CURRICULUM TOPICS (use these as reference for H2 headings):
The following are known medical topics from the curriculum. If a paragraph closely matches one of these, it is almost certainly an H2 heading:
{chr(10).join(f"  - {t}" for t in sample)}

"""

    prompt = f"""You are analyzing a medical study document (PA/physician assistant exam prep).
The document's paragraphs are listed below with formatting clues in brackets.
Your job: identify which paragraphs are HEADINGS and assign each a level 1–4.
{curriculum_section}
Heading level guide:
- H1: Major subject area (e.g. "CARDIOLOGY", "Emergency Medicine", "Surgery") — rare, 0-2 per doc
- H2: Specific condition or topic matching the curriculum topics above (e.g. "Congestive Heart Failure", "COPD") — these become card sections
- H3: Sub-topic within a condition (e.g. "Etiology", "Pathophysiology", "Diagnosis", "Treatment", "Clinical Presentation")
- H4: Further subdivision (e.g. "Lab Values", "Imaging", "Medications")

Additional clues a paragraph is a heading:
- Short text (< 80 chars), bold or larger font, flush left
- ALL CAPS or Title Case with no trailing punctuation
- Known sub-section keywords: Etiology, Pathophysiology, Signs/Symptoms, Diagnosis, Treatment, Prognosis, Complications, Epidemiology

Return ONLY a JSON array — no explanation, no markdown fences.
Each item: {{"index": <paragraph index>, "level": <1|2|3|4>}}
Only include paragraphs that are headings. Omit body text.

Paragraphs:
{paragraph_list}"""

    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    response = client.messages.create(
        model=model,
        max_tokens=2048,
        temperature=0,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = response.content[0].text.strip()
    # Strip markdown fences if present
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.strip()

    try:
        detected = json.loads(raw)
        if not isinstance(detected, list):
            logger.warning("Heading detector returned non-list: %r", raw[:200])
            return []
        result = []
        for item in detected:
            if isinstance(item, dict) and "index" in item and "level" in item:
                level = int(item["level"])
                if 1 <= level <= 4:
                    result.append({"index": int(item["index"]), "level": level})
        logger.info("AI heading detector found %d headings in %d paragraphs", len(result), len(hints))
        return result
    except (json.JSONDecodeError, ValueError) as e:
        logger.error("Failed to parse heading detector response: %s | raw: %r", e, raw[:300])
        return []


def parse_docx_with_ai_headings(
    filepath: str,
    model: str = "claude-haiku-4-5-20251001",
    curriculum_version: str = 'v1',
) -> list[dict]:
    """Parse a .docx, using Claude to detect headings when none are found natively.

    First tries the normal parse. If fewer than 2 H2 headings are found,
    runs AI detection (guided by curriculum topics) and re-parses with the
    detected heading map injected.
    """
    from backend.services.doc_processor import parse_docx

    # Normal parse first
    elements = parse_docx(filepath)
    h2_count = sum(1 for e in elements if e.get("type") == "heading" and e.get("level") == 2)

    if h2_count >= 2:
        logger.info("Document already has %d H2 headings — skipping AI detection", h2_count)
        return elements

    logger.info("Document has %d H2 headings — running AI heading detection", h2_count)
    heading_map = detect_headings_with_ai(filepath, model=model, curriculum_version=curriculum_version)

    if not heading_map:
        logger.warning("AI heading detection returned nothing — using original parse")
        return elements

    # Build index → level lookup
    index_to_level = {h["index"]: h["level"] for h in heading_map}

    # Re-parse with heading injection
    return _parse_docx_with_heading_map(filepath, index_to_level)


def _parse_docx_with_heading_map(filepath: str, index_to_level: dict[int, int]) -> list[dict]:
    """Re-parse docx using a pre-computed {paragraph_index: heading_level} map."""
    import re
    import base64
    import docx

    doc = docx.Document(filepath)
    elements = []
    current_headings = {}

    def build_heading_context(headings: dict) -> Optional[str]:
        if not headings:
            return None
        parts = [f"H{l}: {headings[l]}" for l in sorted(headings.keys())]
        return " > ".join(parts)

    para_index = 0
    for para in doc.paragraphs:
        text = para.text.strip()

        # Extract images regardless
        for run in para.runs:
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            drawings = run._element.findall(f'.//{ns}drawing')
            if drawings:
                for drawing in drawings:
                    for blip in drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed:
                            try:
                                rel = para.part.rels[embed]
                                image_data = rel.target_part.blob
                                ct = rel.target_part.content_type or "image/png"
                                b64 = base64.b64encode(image_data).decode('utf-8')
                                elements.append({
                                    "type": "image",
                                    "text": "",
                                    "html": "",
                                    "data_uri": f"data:{ct};base64,{b64}",
                                    "alt_text": None,
                                    "heading_context": build_heading_context(current_headings),
                                })
                            except (KeyError, AttributeError):
                                pass

        if not text:
            para_index += 1
            continue

        # Determine heading level: AI map takes priority, then Word styles
        level = index_to_level.get(para_index)
        if level is None:
            style_name = (para.style.name or "").lower()
            if "heading" in style_name:
                m = re.search(r'(\d)', style_name)
                if m:
                    level = int(m.group(1))

        if level is not None:
            current_headings[level] = text
            for l in list(current_headings.keys()):
                if l > level:
                    del current_headings[l]
            elements.append({
                "type": "heading",
                "text": text,
                "html": f"<h{level}>{text}</h{level}>",
                "level": level,
                "heading_context": build_heading_context(current_headings),
            })
        else:
            # Build run HTML
            html_parts = []
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.bold:
                    t = f"<b>{t}</b>"
                if run.italic:
                    t = f"<i>{t}</i>"
                if run.underline:
                    t = f"<u>{t}</u>"
                html_parts.append(t)
            html = "".join(html_parts) or text

            # List detection
            ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            style_name = (para.style.name or "").lower()
            is_list = "list" in style_name
            list_type = "ol" if ("number" in style_name or "ordered" in style_name) else "ul"
            indent_level = 0

            if not is_list:
                num_pr = para._element.find(f'.//{ns}numPr')
                if num_pr is not None:
                    num_id_el = num_pr.find(f'{ns}numId')
                    if num_id_el is not None:
                        val = int(num_id_el.get(f'{ns}val', '0'))
                        is_list = val != 0

            if is_list:
                ilvl_el = para._element.find(f'.//{ns}ilvl')
                if ilvl_el is not None:
                    indent_level = int(ilvl_el.get(f'{ns}val', '0'))
                elements.append({
                    "type": "list_item",
                    "text": text,
                    "html": html,
                    "list_type": list_type,
                    "indent_level": indent_level,
                    "heading_context": build_heading_context(current_headings),
                })
            else:
                elements.append({
                    "type": "paragraph",
                    "text": text,
                    "html": f"<p>{html}</p>",
                    "heading_context": build_heading_context(current_headings),
                })

        para_index += 1

    # Tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if rows:
            header = " | ".join(rows[0]) if rows else ""
            text = "\n".join(" | ".join(r) for r in rows)
            elements.append({
                "type": "table",
                "text": text,
                "html": _table_to_html(rows),
                "rows": rows,
                "heading_context": build_heading_context(current_headings),
            })

    return elements


def _table_to_html(rows: list[list[str]]) -> str:
    parts = ["<table>"]
    for i, row in enumerate(rows):
        parts.append("<tr>")
        tag = "th" if i == 0 else "td"
        for cell in row:
            parts.append(f"<{tag}>{cell}</{tag}>")
        parts.append("</tr>")
    parts.append("</table>")
    return "".join(parts)
