"""Regression: Word soft line breaks (Shift+Enter) and paste <br> must survive
into the source text we send the AI, instead of collapsing into one run-together
block. This is the fix for the sibling-split inconsistency (flattened structure
made the model's split decision a coin-flip at temperature 0)."""
from backend.services.generator import (
    markdown_source_from_html, structured_source_from_html, build_generation_prompt,
)


def test_soft_breaks_become_separate_lines():
    # The "Management:" block from the client's screenshot — one <p>, soft breaks.
    html = ("<p>Management:<br>Surgical management — resection"
            "<br>Follow-up — every 3 months<br>Emotional support — counseling</p>")
    out = markdown_source_from_html(html)
    lines = out.split("\n")
    assert lines == [
        "Management:",
        "Surgical management — resection",
        "Follow-up — every 3 months",
        "Emotional support — counseling",
    ]
    # The three sub-items must NOT be on one line (the bug).
    assert "resection Follow-up" not in out


def test_recurrent_abortion_shape():
    html = ("<p>Recurrent abortion<br>Definition: 3+ losses"
            "<br>Causes: genetic, anatomic<br>Evaluation: karyotype, HSG</p>")
    out = markdown_source_from_html(html)
    assert out.split("\n") == [
        "Recurrent abortion",
        "Definition: 3+ losses",
        "Causes: genetic, anatomic",
        "Evaluation: karyotype, HSG",
    ]


def test_real_nested_lists_still_nest():
    # No regression: genuine <ul><li> nesting keeps its indentation.
    html = "<ul><li>Causes<ul><li>Genetic</li><li>Acquired</li></ul></li><li>Evaluation</li></ul>"
    assert markdown_source_from_html(html).split("\n") == [
        "- Causes",
        "  - Genetic",
        "  - Acquired",
        "- Evaluation",
    ]


def test_bold_preserved_across_break():
    html = "<p><b>Migraine</b><br>Aura present<br>Photophobia</p>"
    assert markdown_source_from_html(html).split("\n") == [
        "**Migraine**", "Aura present", "Photophobia",
    ]


def test_margin_indent_nesting_preserved():
    html = ('<p>Top</p><p style="margin-left:1.5em">Child A</p>'
            '<p style="margin-left:1.5em">Child B</p>')
    assert markdown_source_from_html(html).split("\n") == ["Top", "  Child A", "  Child B"]


def test_structured_source_splits_breaks_too():
    html = "<p>Header<br>Line one<br>Line two</p>"
    assert structured_source_from_html(html).split("\n") == ["Header", "Line one", "Line two"]


def test_list_item_with_soft_breaks_keeps_marker_on_first_line_only():
    html = "<ul><li>Point A<br>detail one<br>detail two</li><li>Point B</li></ul>"
    assert markdown_source_from_html(html).split("\n") == [
        "- Point A", "detail one", "detail two", "- Point B",
    ]


def test_connecting_sentence_present_in_user_prompt():
    _, user = build_generation_prompt(
        {"content_html": "<p>x</p>", "heading": "H"}, "RULES",
    )
    assert "applying every rule and the exact output format" in user
    assert "its own sibling card" in user


def test_parse_docx_converts_soft_breaks_to_br():
    # Build a docx with a soft-break paragraph and confirm content_html carries <br>.
    import docx
    from docx import Document
    from backend.services.doc_processor import parse_docx, build_content_html
    d = Document()
    p = d.add_paragraph()
    p.add_run("Management:")
    br = p.add_run(); br._element.append(docx.oxml.OxmlElement("w:br"))
    p.add_run("Surgical management")
    import tempfile, os
    fd, path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        d.save(path)
        html = build_content_html(parse_docx(path))
        assert "<br>" in html
        assert "Management:<br>Surgical management" in html
    finally:
        os.unlink(path)
